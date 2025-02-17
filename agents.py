import multiprocessing
import os

import docx
from crewai import LLM, Agent
from fuzzywuzzy import fuzz
from langchain.tools import tool
from PyPDF2 import PdfReader

llm = LLM(
    model="gemini/gemini-2.0-flash",
    temperature=0.7,
    verbose=True,
    api_key=os.getenv("GOOGLE_API_KEY"),
)
print(f"Using model: {llm.model}")


# @tool
# def extract_blueprint_rules(blueprint_path: str) -> list:
#     """Parse company blueprints (.txt, .docx, .pdf) to extract coding rules."""

#     def read_txt(file_path):
#         with open(file_path, "r") as file:
#             return file.readlines()

#     def read_docx(file_path):
#         doc = docx.Document(file_path)
#         return [p.text for p in doc.paragraphs]

#     def read_pdf(file_path):
#         reader = PdfReader(file_path)
#         content = [
#             page.extract_text() or "[Unable to extract text]" for page in reader.pages
#         ]
#         return content

#     file_ext = os.path.splitext(blueprint_path)[-1].lower()
#     if file_ext == ".txt":
#         content = read_txt(blueprint_path)
#     elif file_ext == ".docx":
#         content = read_docx(blueprint_path)
#     elif file_ext == ".pdf":
#         content = read_pdf(blueprint_path)
#     else:
#         raise ValueError("Unsupported file type. Use .txt, .docx, or .pdf.")

#     rules = [
#         rule.strip() for line in content for rule in line.split("\n") if rule.strip()
#     ]
#     return rules


@tool
def extract_blueprint_rules(blueprint_path: str) -> list:
    """Parse company blueprints (.txt, .docx, .pdf) to extract coding rules."""

    def read_txt(file_path):
        with open(file_path, "r") as file:
            return file.readlines()

    def read_docx(file_path):
        doc = docx.Document(file_path)
        return [p.text for p in doc.paragraphs]

    def read_pdf(file_path):
        reader = PdfReader(file_path)
        content = [
            page.extract_text() or "[Unable to extract text]" for page in reader.pages
        ]
        return content

    file_ext = os.path.splitext(blueprint_path)[-1].lower()
    print(f"Extracting rules from: {blueprint_path} ({file_ext})")

    if file_ext == ".txt":
        content = read_txt(blueprint_path)
    elif file_ext == ".docx":
        content = read_docx(blueprint_path)
    elif file_ext == ".pdf":
        content = read_pdf(blueprint_path)
    else:
        raise ValueError("Unsupported file type. Use .txt, .docx, or .pdf.")

    rules = [
        rule.strip() for line in content for rule in line.split("\n") if rule.strip()
    ]

    return rules


blueprint_parser_agent = Agent(
    role="BluePrint Parser",
    goal="Extract coding guidelines from the files in the blueprints folder",
    tools=[extract_blueprint_rules],
    backstory="Expert in company guidelines and coding practices.",
    llm=llm,
    verbose=True,
)


# @tool
# def check_code(base_path: str) -> dict:
#     """Analyze Python codebase and extract functions and imports."""

#     def extract_components(file_path):
#         with open(file_path, "r") as file:
#             lines = file.readlines()
#         functions = [line.strip() for line in lines if line.strip().startswith("def ")]
#         imports = [
#             line.strip()
#             for line in lines
#             if line.strip().startswith("import") or "from " in line
#         ]
#         return {"functions": functions, "imports": imports}

#     results = {"functions": [], "imports": []}
#     for root, _, files in os.walk(base_path):
#         for file in files:
#             if file.endswith(".py"):
#                 components = extract_components(os.path.join(root, file))
#                 results["functions"].extend(components["functions"])
#                 results["imports"].extend(components["imports"])
#     return results


@tool
def check_code(base_path: str) -> dict:
    """Analyze a Python file or directory and extract functions, imports, and classes using multiprocessing."""

    def extract_components(file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            lines = file.readlines()
        return {
            "functions": [
                line.strip() for line in lines if line.strip().startswith("def ")
            ],
            "imports": [
                line.strip()
                for line in lines
                if line.strip().startswith("import") or line.strip().startswith("from ")
            ],
            "classes": [
                line.strip() for line in lines if line.strip().startswith("class ")
            ],
        }

    results = {"functions": [], "imports": [], "classes": []}

    if os.path.isfile(base_path):
        components = extract_components(base_path)
        results["functions"].extend(components["functions"])
        results["imports"].extend(components["imports"])
        results["classes"].extend(components["classes"])
    else:
        with multiprocessing.Pool() as pool:
            file_paths = [
                os.path.join(root, file)
                for root, _, files in os.walk(base_path)
                for file in files
                if file.endswith(".py")
            ]
            extracted_data = pool.map(extract_components, file_paths)

        for data in extracted_data:
            results["functions"].extend(data["functions"])
            results["imports"].extend(data["imports"])
            results["classes"].extend(data["classes"])

    return results


code_analyzer_agent = Agent(
    role="Code Analyzer",
    goal="Extract key components from the codebase",
    tools=[check_code],
    backstory="Senior Engineer specializing in code analysis.",
    llm=llm,
    verbose=True,
)


@tool
def verify_code_against_blueprint(code_analysis: dict, blueprint_rules: list) -> dict:
    """Compare codebase components with blueprint rules using fuzzy matching."""

    violations = []

    for rule in blueprint_rules:
        found_match = False
        for func in code_analysis["functions"]:
            similarity = fuzz.partial_ratio(rule.lower(), func.lower())
            if similarity > 80:  # Adjust threshold as needed
                found_match = True
                break

        if not found_match:
            violations.append(f"Rule Violation: {rule} not implemented in functions.")

    return {"violations": violations}


code_verifier_agent = Agent(
    role="Code Verifier",
    goal="Ensure the code follows company blueprints",
    tools=[verify_code_against_blueprint],
    backstory="Strict compliance officer enforcing rules.",
    llm=llm,
    verbose=True,
)


@tool
def generate_compliance_report(violations: dict) -> str:
    """Generate a formatted compliance report."""
    report = "🚨 Compliance Report 🚨\n\n"
    if violations["violations"]:
        report += "\n".join(violations["violations"])
    else:
        report += "✅ No violations found. Code is compliant!"
    return report


report_generator_agent = Agent(
    role="Report Generator",
    goal="Generate compliance reports",
    tools=[generate_compliance_report],
    backstory="Documentation expert for compliance.",
    llm=llm,
    verbose=True,
)
